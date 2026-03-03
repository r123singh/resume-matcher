
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

def save_as_pdf(data,filename, resume_dir):
    # Create a PDF
    pdf_filepath = os.path.join(resume_dir, f"{filename}.pdf")
    doc = SimpleDocTemplate(pdf_filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        name="Title",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        spaceAfter=12,
        alignment=1,  # Center alignment
    )
    header_style = ParagraphStyle(
        name="Header",
        parent=styles["Heading2"],
        fontSize=14,
        leading=18,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        name="Body",
        parent=styles["BodyText"],
        fontSize=12,
        leading=16,
        spaceAfter=12,
    )
    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["BodyText"],
        fontSize=12,
        leading=16,
        spaceAfter=6,
        leftIndent=10,
        bulletIndent=0,
        bulletFontName="Helvetica-Bold",
        bulletFontSize=12,
    )

    # Add name as title
    story.append(Paragraph(data["name"], title_style))

    # Add contact information with icons and hyperlinks
    contact_info = (
        f"📍New Delhi | "
        f"📞{data['contact_information']['phone']} | "
        f"📧<a href='mailto:{data['contact_information']['email']}' color='blue'>{data['contact_information']['email']}</a> | "
        f"🔗<a href='{data['contact_information']['linkedIn']}' color='blue'>LinkedIn</a> | "
        f"🔗<a href='{data['contact_information']['github']}' color='blue'>GitHub</a>"
    )
    story.append(Paragraph(contact_info, body_style))
    story.append(Spacer(1, 12))

    # Add summary
    story.append(Paragraph("Summary", header_style))
    story.append(Paragraph(data["summary"], body_style))
    story.append(Spacer(1, 12))

    # Add skills
    story.append(Paragraph("Core Competencies:", header_style))
    for skill in data["skills"]["core_competencies"]:
        story.append(Paragraph(f"• {skill}", bullet_style))
    story.append(Spacer(1, 12))

    # Add work experience
    story.append(Paragraph("Work Experience", header_style))
    for exp in data["work_experience"]:
        story.append(Paragraph(f"{exp['title']} | {exp['company']}", body_style))
        story.append(Paragraph(f"{exp['duration']} | {exp['location']}", body_style))
        story.append(Paragraph(f"{exp['main_description']}"), body_style)
        for desc in exp["points_description"]:
            story.append(Paragraph(f"{desc}", bullet_style))
        story.append(Spacer(1, 6))
    story.append(Spacer(1, 12))

    # Add education
    story.append(Paragraph("Education", header_style))
    for edu in data["education"]:
        story.append(Paragraph(f"{edu['degree_type']} - {edu['institution']} ({edu['year']})", body_style))
    story.append(Spacer(1, 12))

    # Add certifications
    story.append(Paragraph("Certifications", header_style))
    for cert in data["certifications"]:
        story.append(Paragraph(f"{cert}", bullet_style))
    story.append(Spacer(1, 12))
    
    # Add technical skills
    story.append(Paragraph("Technical Skills:", header_style))
    for skill in data["skills"]["technical_skills"]:
        story.append(Paragraph(f"{skill}", bullet_style))
    story.append(Spacer(1, 12))

    # Build the PDF
    doc.build(story)

    print(f"PDF generated successfully: {pdf_filepath}")

    
def save_as_docx(data, filename, resume_dir):
        # Create a Word document
    doc = Document()
    
    doc.styles["Normal"].font.name = "Cambria"
    doc.styles["Normal"].font.size = Pt(11)

    # Add name as title
    title = doc.add_paragraph(data["name"])
    title.style = "Title"
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment
    title.runs[0].font.name = "Cambria"
    title.runs[0].font.size = Pt(30)
    title.runs[0].bold = True

    # Add contact information with icons and hyperlinks
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.add_run(f"📍{data['contact_information']['location']} | ").font.name = "Cambria"
    contact_paragraph.add_run(f"📞{data['contact_information']['phone']} | ").font.name = "Cambria"
    contact_paragraph.add_run(f"📧").font.name = "Cambria"
    contact_paragraph.add_run(f"{data['contact_information']['email']}").font.name = "Cambria"
    contact_paragraph.add_run(" | ").font.name = "Cambria"
    contact_paragraph.add_run("🔗").font.name = "Cambria"
    contact_paragraph.add_run("LinkedIn").font.name = "Cambria"
    contact_paragraph.add_run(" | ").font.name = "Cambria"
    contact_paragraph.add_run("🔗").font.name = "Cambria"
    contact_paragraph.add_run("GitHub").font.name = "Cambria"

    # Add hyperlinks for email, LinkedIn, and GitHub
    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    # Add hyperlinks to email, LinkedIn, and GitHub
    add_hyperlink(contact_paragraph, data["contact_information"]["email"], f"mailto:{data['contact_information']['email']}")
    add_hyperlink(contact_paragraph, "LinkedIn", data["contact_information"]["linkedIn"])
    add_hyperlink(contact_paragraph, "GitHub", data["contact_information"]["github"])

    # Add summary
    doc.add_heading("Summary", level=1)
    summary_paragraph = doc.add_paragraph(data["summary"])
    summary_paragraph.style = "Body Text"
    summary_paragraph.runs[0].font.name = "Cambria"
    

    # Add skills
    doc.add_heading("Skills", level=1)
    doc.add_heading("Core Competencies:", level=2)
    for skill in data["skills"]["core_competencies"]:
        doc.add_paragraph(f"{skill}", style="List Bullet").runs[0].font.name = "Cambria"

    doc.add_heading("Technical Skills:", level=2)
    for skill in data["skills"]["technical_skills"]:
        doc.add_paragraph(f"{skill}", style="List Bullet").runs[0].font.name = "Cambria"

    # Add work experience
    doc.add_heading("Work Experience", level=1)
    for exp in data["work_experience"]:
        doc.add_heading(f"{exp['title']}|{exp['company']}", level=2)
        durlocation = doc.add_paragraph(f"{exp['duration']}|{exp['location']}")
        durlocation.runs[0].italic = True
        durlocation.runs[0].font.name = "Cambria"
        doc.add_paragraph(f"{exp['main_description']}").runs[0].font.name = "Cambria"
        for desc in exp["points_description"]:
            doc.add_paragraph(f"{desc}", style="List Bullet").runs[0].font.name = "Cambria"

    # Add education
    doc.add_heading("Education", level=1)
    for edu in data["education"]:
        doc.add_paragraph(f"{edu['degree_type']} - {edu['institution']} ({edu['year']})").runs[0].font.name = "Cambria"

    # Add certifications
    doc.add_heading("Certifications", level=1)
    for cert in data["certifications"]:
        doc.add_paragraph(f"{cert}", style="List Bullet").runs[0].font.name = "Cambria"

    # Save the document
    docx_filepath = os.path.join(resume_dir, f"{filename}.docx")
    doc.save(docx_filepath)

    print(f"Resume Word document generated successfully: {docx_filepath}")